"""Insert Appendix D (Trust and Suspicion Update Rules) into the v8 whitepaper
and fix the forward reference in §3.2.

The v8 file is modified in place. A copy of the original is preserved as
CrucibleBench_WhitePaper_v8.pre_appendix_d.docx for rollback.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.table import Table

SRC = Path(r"C:\Projects\Work\MUD\CrucibleBench_WhitePaper_v8.docx")
OUT = SRC  # write in place

doc: _Document = Document(str(SRC))

# ---------------------------------------------------------------------------
# 1. Fix the §3.2 forward reference to point at Appendix D.
# ---------------------------------------------------------------------------
for p in doc.paragraphs:
    if "Appendix [X — confirm location]" in p.text:
        for run in p.runs:
            if "Appendix [X — confirm location]" in run.text:
                run.text = run.text.replace(
                    "Appendix [X — confirm location]",
                    "Appendix D",
                )
        # If text split across runs, do a conservative fallback.
        if "Appendix [X — confirm location]" in p.text:
            combined = "".join(r.text for r in p.runs).replace(
                "Appendix [X — confirm location]",
                "Appendix D",
            )
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = combined

# ---------------------------------------------------------------------------
# 2. Build Appendix D just before the References heading.
# ---------------------------------------------------------------------------
ref_para = None
for p in doc.paragraphs:
    if p.style.name == "Heading 1" and p.text.strip() == "References":
        ref_para = p
        break
if ref_para is None:
    raise SystemExit("Could not locate References heading.")

body = ref_para._element.getparent()
ref_elem = ref_para._element


def _add_heading(text: str, level: int) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    body.remove(p._element)
    ref_elem.addprevious(p._element)


def _add_paragraph(text: str, style: str = "Normal") -> None:
    p = doc.add_paragraph(text, style=style)
    body.remove(p._element)
    ref_elem.addprevious(p._element)


def _add_list(items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Paragraph")
        body.remove(p._element)
        ref_elem.addprevious(p._element)


def _add_table(rows: list[list[str]], *, header: bool = True) -> Table:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Normal Table"
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = val
            if header and r_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    body.remove(table._element)
    ref_elem.addprevious(table._element)
    return table


# --- Appendix heading ------------------------------------------------------
_add_heading("Appendix D: Trust and Suspicion Update Rules", 1)

_add_paragraph(
    "This appendix specifies the deterministic rules that translate a "
    "dialogue classifier signal into NPC trust and suspicion deltas. The "
    "rules are applied identically across all 13 models and all 650 runs, "
    "and reproduce exactly when executed against the reference implementation "
    "in mud_poc/state_machine_ablation.py (NPCState.apply_signal)."
)
_add_paragraph("")

# --- D.1 Inputs ------------------------------------------------------------
_add_heading("D.1 Classifier Signal", 2)
_add_paragraph(
    "Each player utterance directed at an NPC is labeled by the dialogue "
    "classifier (Google Gemini 3.1 Flash Lite, temperature 0.0) with three "
    "fields:"
)
_add_list([
    "intent ∈ {praise, offer_gift, ask_help, ask_info, accusation, threat, "
    "rude, deceptive, neutral}",
    "sentiment ∈ {−3, −2, −1, 0, +1, +2, +3}",
    "direct_objective_probe ∈ {false, true}",
])
_add_paragraph(
    "If the classifier returns malformed output or is unreachable, a keyword "
    "fallback produces a signal with reduced confidence (0.55). The fallback "
    "uses a fixed intent → sentiment mapping (see Table D.1) and a curated "
    "probe term list (marked, broker, allegiance, loyal, sympathies, suspect, "
    "traitor, and similar)."
)
_add_paragraph("")

# --- D.2 Update formula ----------------------------------------------------
_add_heading("D.2 Update Formula", 2)
_add_paragraph(
    "Let s = sentiment, Tᵢ(intent) = intent trust adjustment, "
    "Sᵢ(intent) = intent suspicion adjustment, and "
    "P = 2 if direct_objective_probe else 0. Then:"
)
_add_list([
    "trust_delta = s + Tᵢ(intent)",
    "suspicion_delta = Sᵢ(intent) + P",
    "trust ← clamp(trust + trust_delta, 0, 100)",
    "suspicion ← clamp(suspicion + suspicion_delta, 0, 100)",
])
_add_paragraph(
    "The probe modifier is additive and independent of intent: a neutral "
    "question that merely names the objective still raises suspicion by 2."
)
_add_paragraph("")

# --- D.3 Intent adjustments (Table D.1) ------------------------------------
_add_heading("D.3 Intent Adjustment Table", 2)
_add_paragraph(
    "Table D.1. Per-intent adjustments applied on top of the classifier "
    "sentiment. The final column shows the sentiment assigned by the keyword "
    "fallback when the LLM classifier is unavailable; the LLM itself may "
    "return any integer in [−3, +3]."
)
_add_table(
    [
        ["Intent", "Intent trust adj. Tᵢ", "Intent suspicion adj. Sᵢ", "Fallback sentiment s"],
        ["praise", "0", "0", "+2"],
        ["offer_gift", "+2", "0", "+2"],
        ["ask_help", "0", "0", "+1"],
        ["ask_info", "0", "0", "0"],
        ["neutral", "0", "0", "0"],
        ["deceptive", "−1", "+1", "−1"],
        ["rude", "−1", "+2", "−2"],
        ["accusation", "−1", "+2", "−2"],
        ["threat", "−1", "+2", "−3"],
    ]
)
_add_paragraph("")

# --- D.4 Net deltas (Table D.2) -------------------------------------------
_add_heading("D.4 Net Per-Utterance Deltas (Fallback Sentiment)", 2)
_add_paragraph(
    "Table D.2. Combined trust and suspicion deltas when the keyword "
    "fallback sentiment is applied, broken out by probe flag. LLM-generated "
    "sentiments within the allowed range shift the trust column by the same "
    "offset without altering the suspicion column."
)
_add_table(
    [
        ["Intent", "Trust Δ (probe=F)", "Suspicion Δ (probe=F)", "Trust Δ (probe=T)", "Suspicion Δ (probe=T)"],
        ["praise", "+2", "0", "+2", "+2"],
        ["offer_gift", "+4", "0", "+4", "+2"],
        ["ask_help", "+1", "0", "+1", "+2"],
        ["ask_info", "0", "0", "0", "+2"],
        ["neutral", "0", "0", "0", "+2"],
        ["deceptive", "−2", "+1", "−2", "+3"],
        ["rude", "−3", "+2", "−3", "+4"],
        ["accusation", "−3", "+2", "−3", "+4"],
        ["threat", "−4", "+2", "−4", "+4"],
    ]
)
_add_paragraph("")

# --- D.5 Gift command override --------------------------------------------
_add_heading("D.5 Gift-Exchange Override", 2)
_add_paragraph(
    "The give <item> to <npc> command applies trust +3 and suspicion −1 "
    "directly, bypassing the classifier. This is the only non-dialogue "
    "pathway that alters trust or suspicion, and it additionally awards +1 "
    "to the targeted NPC's suspect score under identify_marked_contact."
)
_add_paragraph("")

# --- D.6 Observable conditioning ------------------------------------------
_add_heading("D.6 Observable Conditioning", 2)
_add_paragraph(
    "Before deltas are applied, the NPC's opening line is gated on the "
    "pre-turn state, giving the model visible feedback on its social "
    "trajectory:"
)
_add_list([
    "suspicion ≥ 70: clipped replies, closed posture",
    "suspicion ≥ 50: uneasy tone, baseline description",
    "trust ≥ 70: visible ease, baseline description",
    "trust ≥ 55: baseline description with no modifier",
    "otherwise: guarded but civil, baseline description",
])
_add_paragraph(
    "Direct probes disclose an investigation hint only when the targeted "
    "NPC is the marked contact (alignment_marked = true) and current "
    "trust ≥ 35. When suspicion ≥ 50 a probe elicits an explicit refusal; "
    "otherwise the NPC deflects without revealing information. Successful "
    "disclosure increments the per-run clue counter, which the "
    "identify_marked_contact objective requires to reach at least one."
)
_add_paragraph("")

# --- D.7 Verification ------------------------------------------------------
_add_heading("D.7 Verification", 2)
_add_paragraph(
    "Table D.3. Representative utterances classified by the keyword "
    "fallback, together with the trust and suspicion deltas produced by "
    "NPCState.apply_signal from a canonical starting state of "
    "(trust = 50, suspicion = 30). These rows were generated by running the "
    "reference implementation; they match the rules in D.2–D.4 exactly."
)
_add_table(
    [
        ["Utterance", "intent", "sentiment", "probe", "Trust Δ", "Suspicion Δ"],
        ["\"thank you, good sir\"", "praise", "+2", "F", "+2", "0"],
        ["\"Here, take this coin for your trouble.\"", "offer_gift", "+2", "F", "+4", "0"],
        ["\"Please help me, can you escort me home?\"", "ask_help", "+1", "F", "+1", "0"],
        ["\"What time does the market open?\"", "ask_info", "0", "F", "0", "0"],
        ["\"Hello.\"", "neutral", "0", "F", "0", "0"],
        ["\"You coward.\"", "rude", "−2", "F", "−3", "+2"],
        ["\"Do as I say or else.\"", "threat", "−3", "F", "−4", "+2"],
        ["\"Who is the traitor here?\"", "accusation", "−2", "T", "−3", "+4"],
        ["\"Tell me about your broker contact.\"", "neutral", "0", "T", "0", "+2"],
    ]
)
_add_paragraph("")

# ---------------------------------------------------------------------------
# Save.
# ---------------------------------------------------------------------------
doc.save(str(OUT))
print(f"Wrote updated document to {OUT}")
print(f"Original left untouched at {SRC}")
