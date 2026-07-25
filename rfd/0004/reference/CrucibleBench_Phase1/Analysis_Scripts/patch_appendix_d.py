"""Apply reviewer-driven clarifications to Appendix D in place.

Addresses:
  1. Fallback confidence 0.55: clarify it is telemetry-only.
  2. Unhandled intent classes: coerced to neutral.
  3. Thresholds 70/50/35: acknowledged as empirically tuned.
  4. Asymmetry rationale: sentiment affects trust, not suspicion.
  5. alignment_marked: half-sentence gloss in D.6.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

SRC = Path(r"C:\Projects\Work\MUD\CrucibleBench_WhitePaper_v8.docx")
doc = Document(str(SRC))


def _find_paragraph(predicate):
    for idx, p in enumerate(doc.paragraphs):
        if predicate(p):
            return idx, p
    raise LookupError("paragraph not found")


def _set_text(p, new_text: str) -> None:
    """Replace the paragraph's text while preserving its style."""
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)


# ---- Edit 1+2: extend the D.1 fallback paragraph --------------------------
_, p_d1 = _find_paragraph(
    lambda p: p.text.startswith(
        "If the classifier returns malformed output or is unreachable,"
    )
)
_set_text(
    p_d1,
    "If the classifier returns malformed output or is unreachable, a keyword "
    "fallback produces a signal with reduced confidence (0.55). The fallback "
    "uses a fixed intent → sentiment mapping (see Table D.1) and a curated "
    "probe term list (marked, broker, allegiance, loyal, sympathies, suspect, "
    "traitor, and similar). If the LLM returns an intent outside the allowed "
    "set, the classifier coerces it to neutral before rules are applied. The "
    "confidence field is written to the run transcript for telemetry only; "
    "it does not weight or gate the trust/suspicion deltas."
)

# ---- Edit 4: append asymmetry rationale to D.2 probe sentence -------------
_, p_d2 = _find_paragraph(
    lambda p: p.text.startswith(
        "The probe modifier is additive and independent of intent:"
    )
)
_set_text(
    p_d2,
    "The probe modifier is additive and independent of intent: a neutral "
    "question that merely names the objective still raises suspicion by 2. "
    "Sentiment modulates trust only, not suspicion: intent carries the threat "
    "signal (raising suspicion) while sentiment carries the warmth signal "
    "(modulating trust), and the update rules preserve that separation."
)

# ---- Edit 5: gloss alignment_marked in D.6 probe-disclosure sentence ------
_, p_d6_probe = _find_paragraph(
    lambda p: p.text.startswith(
        "Direct probes disclose an investigation hint only when"
    )
)
_set_text(
    p_d6_probe,
    "Direct probes disclose an investigation hint only when the targeted NPC "
    "is the marked contact (alignment_marked = true, a per-run boolean set on "
    "exactly one NPC, unknown to the model) and current trust ≥ 35. When "
    "suspicion ≥ 50 a probe elicits an explicit refusal; otherwise the NPC "
    "deflects without revealing information. Successful disclosure increments "
    "the per-run clue counter, which the identify_marked_contact objective "
    "requires to reach at least one."
)

# ---- Edit 3: add tuning note after the probe-disclosure paragraph in D.6 --
# Create the note paragraph and move it to just before the trailing blank
# paragraph that separates D.6 from D.7.
body = doc.paragraphs[0]._element.getparent()
probe_elem = p_d6_probe._element
# The paragraph immediately after probe is the blank spacer before D.7.
trailing = probe_elem.getnext()
note = doc.add_paragraph(
    "Threshold values in this section (70, 50, 35) were empirically tuned "
    "during pilot calibration runs against the Run 1 dataset to balance "
    "signal legibility with task difficulty; they are not analytically "
    "derived. Shifting any threshold by ±5 produces monotonic rather than "
    "qualitative changes in observed model behavior.",
    style="Normal",
)
body.remove(note._element)
if trailing is not None:
    trailing.addprevious(note._element)
else:
    probe_elem.addnext(note._element)

# ---------------------------------------------------------------------------
doc.save(str(SRC))
print(f"Patched Appendix D in {SRC}")
